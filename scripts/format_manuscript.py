#!/usr/bin/env python3
"""Format a manuscript DOCX for Nature-family submission workflows.

This script is intentionally conservative about scientific content. It applies a
deterministic Word layout, maps recognizable manuscript elements to named
styles, adds continuous line numbers and page numbers, formats tables, and
writes a JSON audit. It does not invent, rewrite, or silently reorder content.

Supported journal profiles:
  - nature-communications: deterministic visual tokens derived from the AJE
    Nature Communications submission template changed in December 2025. An
    optional local copy can be supplied with --template.
  - nature: readable initial-submission formatting plus official structure audit.
  - nature-methods: the same visual formatting as Nature plus a separate
    Nature Methods article-type and submission-stage audit.

Examples:
  python format_manuscript.py draft.docx --journal nature-communications \
      --output draft_ncomms.docx --report draft_ncomms.json
  python format_manuscript.py draft.docx --journal nature-methods \
      --article-type article --stage initial --output draft_nmeth.docx
  python format_manuscript.py draft.docx --journal nature --audit-only \
      --report nature_audit.json
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable, Sequence

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
DEFAULT_NCOMMS_TEMPLATE = (
    SKILL_DIR
    / "assets"
    / "AJE-Nature-Communications-Submission-Template_changed_Dec2025.docx"
)


@dataclass(frozen=True)
class Profile:
    name: str
    visual_format: str
    title_limit_kind: str
    title_limit: int
    abstract_limit: int | None
    display_limit: int | None
    required_headings: tuple[str, ...]
    recommended_headings: tuple[str, ...]
    ordered_headings: tuple[str, ...]
    introduction_without_heading: bool = False
    discussion_without_subheadings: bool = False


PROFILES: dict[tuple[str, str], Profile] = {
    ("nature-communications", "article"): Profile(
        name="nature-communications",
        visual_format="nature-communications",
        title_limit_kind="words",
        title_limit=15,
        abstract_limit=150,
        display_limit=10,
        required_headings=("abstract", "introduction", "results", "methods", "references"),
        recommended_headings=(
            "discussion",
            "data availability",
            "code availability",
            "acknowledgements",
            "author contributions",
            "competing interests",
        ),
        ordered_headings=(
            "abstract",
            "introduction",
            "results",
            "discussion",
            "methods",
            "data availability",
            "code availability",
            "references",
            "acknowledgements",
            "funding",
            "author contributions",
            "competing interests",
            "supplementary information",
        ),
        discussion_without_subheadings=True,
    ),
    ("nature", "article"): Profile(
        name="nature",
        visual_format="nature",
        title_limit_kind="characters",
        title_limit=75,
        abstract_limit=200,
        display_limit=6,
        required_headings=("references", "methods"),
        recommended_headings=(
            "data availability",
            "code availability",
            "acknowledgements",
            "funding",
            "author contributions",
            "competing interests",
        ),
        ordered_headings=(
            "references",
            "figure legends",
            "methods",
            "methods references",
            "data availability",
            "code availability",
            "acknowledgements",
            "funding",
            "author contributions",
            "competing interests",
            "additional information",
            "extended data",
        ),
    ),
    ("nature-methods", "article"): Profile(
        name="nature-methods",
        visual_format="nature",
        title_limit_kind="none",
        title_limit=0,
        abstract_limit=150,
        display_limit=6,
        required_headings=("results", "discussion", "methods", "references"),
        recommended_headings=(
            "data availability",
            "code availability",
            "acknowledgements",
            "funding",
            "author contributions",
            "competing interests",
        ),
        ordered_headings=(
            "results",
            "discussion",
            "references",
            "methods",
            "data availability",
            "code availability",
            "methods references",
            "acknowledgements",
            "funding",
            "author contributions",
            "competing interests",
            "extended data",
        ),
        introduction_without_heading=True,
        discussion_without_subheadings=True,
    ),
    ("nature-methods", "resource"): Profile(
        name="nature-methods",
        visual_format="nature",
        title_limit_kind="none",
        title_limit=0,
        abstract_limit=150,
        display_limit=6,
        required_headings=("results", "discussion", "methods", "references"),
        recommended_headings=("data availability", "code availability"),
        ordered_headings=("results", "discussion", "references", "methods", "data availability", "code availability", "methods references"),
        introduction_without_heading=True,
        discussion_without_subheadings=True,
    ),
    ("nature-methods", "analysis"): Profile(
        name="nature-methods",
        visual_format="nature",
        title_limit_kind="none",
        title_limit=0,
        abstract_limit=150,
        display_limit=6,
        required_headings=("results", "discussion", "methods", "references"),
        recommended_headings=("data availability", "code availability"),
        ordered_headings=("results", "discussion", "references", "methods", "data availability", "code availability", "methods references"),
        introduction_without_heading=True,
        discussion_without_subheadings=True,
    ),
    ("nature-methods", "brief-communication"): Profile(
        name="nature-methods",
        visual_format="nature",
        title_limit_kind="none",
        title_limit=0,
        abstract_limit=70,
        display_limit=2,
        required_headings=("methods", "references"),
        recommended_headings=("data availability", "code availability"),
        ordered_headings=("references", "methods", "data availability", "code availability", "methods references"),
        introduction_without_heading=True,
    ),
}


CANONICAL_HEADINGS = {
    "abstract": "abstract",
    "summary": "abstract",
    "introduction": "introduction",
    "results": "results",
    "discussion": "discussion",
    "methods": "methods",
    "online methods": "methods",
    "references": "references",
    "main references": "references",
    "methods references": "methods references",
    "figure legends": "figure legends",
    "figures": "figure legends",
    "data availability": "data availability",
    "data availability statement": "data availability",
    "code availability": "code availability",
    "code availability statement": "code availability",
    "acknowledgements": "acknowledgements",
    "acknowledgments": "acknowledgements",
    "funding": "funding",
    "funding statement": "funding",
    "author contributions": "author contributions",
    "author contributions statement": "author contributions",
    "contributions": "author contributions",
    "competing interests": "competing interests",
    "competing interest": "competing interests",
    "ethics declarations": "ethics declarations",
    "additional information": "additional information",
    "supplementary information": "supplementary information",
    "extended data": "extended data",
    "corresponding author": "corresponding author",
    "affiliations": "affiliations",
}

MAJOR_HEADINGS = set(CANONICAL_HEADINGS.values())
CAPTION_RE = re.compile(
    r"^(?:fig(?:ure)?\.?\s*\d+|extended\s+data\s+fig(?:ure)?\.?\s*\d+|table\s*\d+)\b",
    re.IGNORECASE,
)
WORD_RE = re.compile(r"\b[\w'\-]+\b", re.UNICODE)


def normalize_heading(text: str) -> str:
    text = text.replace("\u00a0", " ").strip().lower()
    text = re.sub(r"^\s*(?:section\s+)?\d+(?:\.\d+)*[.)]?\s+", "", text)
    text = re.sub(r"[:.\s]+$", "", text)
    text = re.sub(r"\s+", " ", text)
    return text


def canonical_heading(text: str) -> str | None:
    return CANONICAL_HEADINGS.get(normalize_heading(text))


def count_words(text: str) -> int:
    return len(WORD_RE.findall(text))


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def validate_docx(path: Path) -> None:
    if not path.is_file():
        raise FileNotFoundError(path)
    if path.suffix.lower() != ".docx":
        raise ValueError(f"Expected a .docx file: {path}")
    if not zipfile.is_zipfile(path):
        raise ValueError(f"Not a valid OOXML package: {path}")
    with zipfile.ZipFile(path) as package:
        required = {"[Content_Types].xml", "word/document.xml"}
        missing = required.difference(package.namelist())
        if missing:
            raise ValueError(f"DOCX is missing required parts: {sorted(missing)}")


def get_profile(journal: str, article_type: str) -> Profile:
    key = (journal, article_type)
    if key not in PROFILES:
        supported = sorted(f"{j}/{a}" for j, a in PROFILES)
        raise ValueError(
            f"Unsupported journal/article type: {journal}/{article_type}. "
            f"Supported: {', '.join(supported)}"
        )
    return PROFILES[key]


def set_rfonts(rpr: Any, font_name: str) -> None:
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def set_half_points(rpr: Any, points: float) -> None:
    value = str(int(round(points * 2)))
    for tag in ("w:sz", "w:szCs"):
        node = rpr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            rpr.append(node)
        node.set(qn("w:val"), value)


def ensure_on_off(rpr: Any, tag: str, enabled: bool) -> None:
    node = rpr.find(qn(tag))
    if node is None:
        node = OxmlElement(tag)
        rpr.append(node)
    node.set(qn("w:val"), "1" if enabled else "0")


def style_font(style: Any, name: str, size: float, bold: bool | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    set_rfonts(rpr, name)
    set_half_points(rpr, size)


def set_style_bottom_border(style: Any, enabled: bool) -> None:
    ppr = style.element.get_or_add_pPr()
    existing = ppr.find(qn("w:pBdr"))
    if existing is not None:
        ppr.remove(existing)
    if not enabled:
        return
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    ppr.append(borders)


def ensure_style(doc: DocumentObject, name: str, *, size: float, bold: bool,
                 line_spacing: float, before: float, after: float,
                 keep_with_next: bool = False, border: bool = False) -> Any:
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles["Normal"]
    style_font(style, "Times New Roman", size, bold)
    pf = style.paragraph_format
    pf.left_indent = None
    pf.right_indent = None
    pf.first_line_indent = None
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing
    pf.keep_with_next = keep_with_next
    set_style_bottom_border(style, border)
    return style


def configure_styles(doc: DocumentObject, profile: Profile) -> dict[str, Any]:
    normal = doc.styles["Normal"]
    style_font(normal, "Times New Roman", 12, False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 2.0

    return {
        "body": ensure_style(
            doc, "BLM Body", size=12, bold=False, line_spacing=2.0,
            before=0, after=0,
        ),
        "title": ensure_style(
            doc, "BLM Title", size=16, bold=True, line_spacing=1.5,
            before=0, after=0, keep_with_next=True,
        ),
        "front": ensure_style(
            doc, "BLM Front Matter", size=12, bold=False, line_spacing=2.0,
            before=0, after=0,
        ),
        "section": ensure_style(
            doc, "BLM Section Heading", size=14, bold=True, line_spacing=1.0,
            before=12, after=0, keep_with_next=True,
            border=profile.visual_format == "nature-communications",
        ),
        "subheading": ensure_style(
            doc, "BLM Subheading", size=14, bold=False, line_spacing=2.0,
            before=12, after=0, keep_with_next=True,
        ),
        "method_subheading": ensure_style(
            doc, "BLM Methods Subheading", size=12, bold=True, line_spacing=2.0,
            before=12, after=0, keep_with_next=True,
        ),
        "caption": ensure_style(
            doc, "BLM Figure or Table Legend", size=11, bold=False,
            line_spacing=2.0, before=12, after=6, keep_with_next=True,
        ),
        "reference": ensure_style(
            doc, "BLM Reference", size=11, bold=False, line_spacing=2.0,
            before=0, after=0,
        ),
        "table": ensure_style(
            doc, "BLM Table Text", size=11, bold=False, line_spacing=2.0,
            before=0, after=0,
        ),
    }


def clear_direct_paragraph_format(paragraph: Any) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    removable = {
        qn("w:spacing"), qn("w:ind"), qn("w:jc"), qn("w:keepNext"),
        qn("w:keepLines"), qn("w:pageBreakBefore"), qn("w:widowControl"),
        qn("w:pBdr"),
    }
    for child in list(ppr):
        if child.tag in removable:
            ppr.remove(child)


def normalize_runs(paragraph: Any, size: float, *, force_bold: bool | None = None) -> None:
    for run_node in paragraph._p.iter(qn("w:r")):
        rpr = run_node.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run_node.insert(0, rpr)
        set_rfonts(rpr, "Times New Roman")
        set_half_points(rpr, size)
        hyperlink = any(parent.tag == qn("w:hyperlink") for parent in run_node.iterancestors())
        if not hyperlink:
            color = rpr.find(qn("w:color"))
            if color is not None:
                rpr.remove(color)
        if force_bold is not None:
            ensure_on_off(rpr, "w:b", force_bold)
            ensure_on_off(rpr, "w:bCs", force_bold)


def set_paragraph_role(paragraph: Any, role: str, styles: dict[str, Any]) -> None:
    clear_direct_paragraph_format(paragraph)
    paragraph.style = styles[role]
    size = {
        "title": 16,
        "section": 14,
        "subheading": 14,
        "method_subheading": 12,
        "caption": 11,
        "reference": 11,
        "front": 12,
        "body": 12,
    }[role]
    force_bold = True if role in {"title", "section", "method_subheading"} else None
    normalize_runs(paragraph, size, force_bold=force_bold)


def classify_paragraphs(paragraphs: Sequence[Any]) -> list[dict[str, Any]]:
    """Classify top-level paragraphs without rebuilding their wrapper list.

    ``python-docx`` reconstructs the complete ``Document.paragraphs`` list on
    every property access. Accepting a cached sequence keeps classification and
    later formatting linear in the number of paragraphs.
    """
    nonempty = [i for i, p in enumerate(paragraphs) if p.text.strip()]
    title_index = None
    for i in nonempty[:15]:
        paragraph = paragraphs[i]
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        if "title" in style_name:
            title_index = i
            break
    if title_index is None and nonempty:
        for i in nonempty[:15]:
            text = paragraphs[i].text.strip()
            if canonical_heading(text) is None and 4 <= count_words(text) <= 40:
                title_index = i
                break
    if title_index is None and nonempty:
        title_index = nonempty[0]

    first_content_heading = len(paragraphs)
    for i, paragraph in enumerate(paragraphs):
        if i <= (title_index if title_index is not None else -1):
            continue
        if canonical_heading(paragraph.text) in {"abstract", "introduction", "results", "methods"}:
            first_content_heading = i
            break

    records: list[dict[str, Any]] = []
    current_section: str | None = None
    in_references = False
    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        heading = canonical_heading(text)
        style_name = paragraph.style.name.lower() if paragraph.style else ""

        if not text and not paragraph._p.xpath(".//w:drawing | .//w:pict | .//m:oMath"):
            records.append({"index": i, "role": "body", "heading": None, "text": text})
            continue
        if i == title_index:
            role = "title"
        elif heading in {"affiliations", "corresponding author", "author contributions"} and i < first_content_heading:
            # Nature Communications places these front-matter labels below the
            # author block as subordinate labels, not ruled main-text sections.
            role = "subheading"
            current_section = heading
            in_references = False
        elif heading == "competing interests" and current_section == "ethics declarations":
            # In the supplied template, Competing interests is nested below the
            # broader Ethics declarations heading.
            role = "subheading"
            in_references = False
        elif heading in MAJOR_HEADINGS:
            role = "section"
            current_section = heading
            in_references = heading in {"references", "methods references"}
        elif CAPTION_RE.match(text):
            role = "caption"
        elif in_references:
            role = "reference"
        elif i < first_content_heading:
            role = "front"
        elif "heading" in style_name or "subheading" in style_name:
            role = "method_subheading" if current_section == "methods" else "subheading"
        else:
            role = "body"
        records.append({"index": i, "role": role, "heading": heading, "text": text})
    return records


def add_line_numbers(section: Any) -> None:
    sectpr = section._sectPr
    existing = sectpr.find(qn("w:lnNumType"))
    if existing is None:
        existing = OxmlElement("w:lnNumType")
        cols = sectpr.find(qn("w:cols"))
        if cols is None:
            sectpr.append(existing)
        else:
            cols.addprevious(existing)
    existing.set(qn("w:countBy"), "1")
    existing.set(qn("w:restart"), "continuous")


def clear_story(story: Any) -> None:
    element = story._element
    for child in list(element):
        element.remove(child)
    element.append(OxmlElement("w:p"))


def suppress_line_number(paragraph: Any) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.find(qn("w:suppressLineNumbers")) is None:
        ppr.append(OxmlElement("w:suppressLineNumbers"))


def add_page_number(paragraph: Any) -> None:
    suppress_line_number(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    normalize_runs(paragraph, 10)


def apply_page_layout(doc: DocumentObject, template_path: Path | None,
                      preserve_header_footer: bool) -> None:
    template_section = None
    if template_path is not None:
        validate_docx(template_path)
        template_section = Document(template_path).sections[0]

    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        if template_section is not None:
            section.page_width = template_section.page_width
            section.page_height = template_section.page_height
            section.top_margin = template_section.top_margin
            section.bottom_margin = template_section.bottom_margin
            section.left_margin = template_section.left_margin
            section.right_margin = template_section.right_margin
            section.header_distance = template_section.header_distance
            section.footer_distance = template_section.footer_distance
        else:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)
        add_line_numbers(section)
        if not preserve_header_footer:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            clear_story(section.header)
            clear_story(section.footer)
            suppress_line_number(section.header.paragraphs[0])
            add_page_number(section.footer.paragraphs[0])

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_cell_shading(cell: Any, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def format_tables(doc: DocumentObject, table_style: Any) -> int:
    count = 0
    for table in doc.tables:
        count += 1
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        table.autofit = True
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if row_index == 0:
                    set_cell_shading(cell, "F2F2F2")
                for paragraph in cell.paragraphs:
                    clear_direct_paragraph_format(paragraph)
                    paragraph.style = table_style
                    normalize_runs(paragraph, 11, force_bold=True if row_index == 0 else None)
        # Mark the first row as repeating when the table spans pages.
        if table.rows:
            trpr = table.rows[0]._tr.get_or_add_trPr()
            if trpr.find(qn("w:tblHeader")) is None:
                trpr.append(OxmlElement("w:tblHeader"))
    return count


def paragraph_text_between(records: list[dict[str, Any]], heading: str) -> str:
    start = None
    pieces: list[str] = []
    for record in records:
        if record["heading"] == heading:
            start = record["index"]
            continue
        if start is not None and record["index"] > start:
            if record["role"] == "section":
                break
            pieces.append(record["text"])
    return " ".join(pieces)


def audit(doc: DocumentObject, records: list[dict[str, Any]], profile: Profile,
          journal: str, article_type: str, stage: str) -> dict[str, Any]:
    headings = [r["heading"] for r in records if r["heading"]]
    title_record = next((r for r in records if r["role"] == "title"), None)
    title = title_record["text"] if title_record else ""
    issues: list[dict[str, str]] = []

    def issue(level: str, code: str, message: str) -> None:
        issues.append({"level": level, "code": code, "message": message})

    for heading in profile.required_headings:
        if heading not in headings:
            issue("Major", "MISSING_REQUIRED_SECTION", f"Missing recognizable section: {heading}.")
    for heading in profile.recommended_headings:
        if heading not in headings:
            issue("Minor", "SECTION_NOT_DETECTED", f"Section not detected; confirm whether applicable: {heading}.")

    # Front-matter labels (for example, Contributions in the supplied Nature
    # Communications template) and nested declaration labels are real headings
    # but are not top-level manuscript sections for order validation.
    seen_order = [
        r["heading"]
        for r in records
        if r["role"] == "section" and r["heading"] in profile.ordered_headings
    ]
    order_values = [profile.ordered_headings.index(h) for h in seen_order]
    if order_values != sorted(order_values):
        issue(
            "Major",
            "SECTION_ORDER",
            "Recognized top-level sections are not in the target profile's expected order. "
            "The formatter does not silently reorder scientific content.",
        )

    if profile.title_limit_kind == "words" and count_words(title) > profile.title_limit:
        issue("Major", "TITLE_LENGTH", f"Title has {count_words(title)} words; target is at most {profile.title_limit}.")
    if profile.title_limit_kind == "characters" and len(title) > profile.title_limit:
        issue("Major", "TITLE_LENGTH", f"Title has {len(title)} characters including spaces; target is at most {profile.title_limit}.")

    abstract_text = paragraph_text_between(records, "abstract")
    if abstract_text and profile.abstract_limit is not None:
        abstract_words = count_words(abstract_text)
        if abstract_words > profile.abstract_limit:
            issue("Major", "ABSTRACT_LENGTH", f"Abstract/summary has {abstract_words} words; target is at most {profile.abstract_limit}.")

    display_count = sum(1 for r in records if r["role"] == "caption")
    if profile.display_limit is not None and display_count > profile.display_limit:
        issue("Major", "DISPLAY_ITEM_LIMIT", f"Detected {display_count} display-item legends; target limit is {profile.display_limit}.")

    if profile.introduction_without_heading and "introduction" in headings:
        issue("Minor", "INTRODUCTION_HEADING", "This article type specifies an Introduction without a visible heading.")

    if article_type == "brief-communication":
        forbidden = [h for h in headings if h in {"introduction", "results", "discussion"}]
        if forbidden:
            issue("Major", "BRIEF_COMMUNICATION_HEADINGS", "Nature Methods Brief Communications should not contain main-text sections or subheadings.")

    if profile.discussion_without_subheadings and "discussion" in headings:
        inside = False
        for record in records:
            if record["heading"] == "discussion":
                inside = True
                continue
            if inside and record["role"] == "section":
                break
            if inside and record["role"] in {"subheading", "method_subheading"}:
                issue("Major", "DISCUSSION_SUBHEADING", "Discussion contains a detected subheading, which this profile does not permit.")
                break

    if journal == "nature-methods" and stage == "initial":
        issue(
            "Style",
            "INITIAL_FORMATTING_FLEXIBLE",
            "Nature Methods does not require special visual formatting at initial submission; "
            "at the user's request, this output uses the Nature visual format while retaining "
            "Nature Methods compliance checks.",
        )
    if journal == "nature" and stage == "initial":
        issue(
            "Style",
            "INITIAL_FORMATTING_FLEXIBLE",
            "Nature is flexible about initial-submission format; this layout applies its recommended 12-point, double-spaced, line-numbered presentation.",
        )

    return {
        "journal": journal,
        "visual_format": profile.visual_format,
        "article_type": article_type,
        "stage": stage,
        "title": title,
        "title_words": count_words(title),
        "title_characters_including_spaces": len(title),
        "abstract_words": count_words(abstract_text) if abstract_text else None,
        "detected_display_item_legends": display_count,
        "detected_headings": headings,
        "issues": issues,
    }


def format_document(input_path: Path, output_path: Path, report_path: Path | None,
                    profile: Profile, journal: str, article_type: str, stage: str,
                    template_path: Path | None, preserve_header_footer: bool,
                    audit_only: bool) -> dict[str, Any]:
    validate_docx(input_path)
    doc = Document(input_path)
    paragraphs = doc.paragraphs
    records = classify_paragraphs(paragraphs)
    report = audit(doc, records, profile, journal, article_type, stage)
    report["input"] = str(input_path)
    report["input_sha256"] = sha256(input_path)

    if not audit_only:
        if input_path.resolve() == output_path.resolve():
            raise ValueError("Refusing to overwrite the input manuscript; choose a different --output path.")
        styles = configure_styles(doc, profile)
        for record in records:
            set_paragraph_role(paragraphs[record["index"]], record["role"], styles)
        selected_template = template_path if journal == "nature-communications" else None
        apply_page_layout(doc, selected_template, preserve_header_footer)
        table_count = format_tables(doc, styles["table"])
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        validate_docx(output_path)
        report["output"] = str(output_path)
        report["output_sha256"] = sha256(output_path)
        report["formatted_paragraphs"] = len(records)
        report["formatted_tables"] = table_count
        if selected_template is not None:
            report["template"] = str(selected_template)
            report["template_sha256"] = sha256(selected_template)

    if report_path is not None:
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    return report


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="Source manuscript DOCX.")
    parser.add_argument(
        "--journal", required=True,
        choices=("nature-communications", "nature", "nature-methods"),
    )
    parser.add_argument(
        "--article-type", default="article",
        choices=("article", "resource", "analysis", "brief-communication"),
    )
    parser.add_argument("--stage", default="initial", choices=("initial", "aip"))
    parser.add_argument("--output", type=Path, help="Formatted DOCX path.")
    parser.add_argument("--report", type=Path, help="Optional JSON audit path.")
    parser.add_argument(
        "--template", type=Path,
        help="Nature Communications template override. Ignored for other journals.",
    )
    parser.add_argument(
        "--preserve-header-footer", action="store_true",
        help="Keep existing headers and footers instead of replacing them with a clean footer page number.",
    )
    parser.add_argument(
        "--audit-only", action="store_true",
        help="Write/print the audit without changing the DOCX.",
    )
    return parser


def main(argv: Iterable[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    if not args.audit_only and args.output is None:
        raise SystemExit("--output is required unless --audit-only is used")
    if args.journal != "nature-methods" and args.article_type != "article":
        raise SystemExit("Only --article-type article is supported for this journal profile")
    profile = get_profile(args.journal, args.article_type)
    template = args.template
    if (
        args.journal == "nature-communications"
        and template is None
        and DEFAULT_NCOMMS_TEMPLATE.is_file()
    ):
        template = DEFAULT_NCOMMS_TEMPLATE
    if template is not None and not template.is_file():
        raise SystemExit(f"Template not found: {template}")

    try:
        report = format_document(
            input_path=args.input,
            output_path=args.output or Path("unused.docx"),
            report_path=args.report,
            profile=profile,
            journal=args.journal,
            article_type=args.article_type,
            stage=args.stage,
            template_path=template,
            preserve_header_footer=args.preserve_header_footer,
            audit_only=args.audit_only,
        )
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2

    print(json.dumps(report, indent=2, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
