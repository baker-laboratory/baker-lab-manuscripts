from __future__ import annotations

import importlib.util
import json
import sys
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from docx import Document
from docx.document import Document as DocumentObject


SKILL_ROOT = Path(__file__).resolve().parents[1]
FORMATTER_PATH = SKILL_ROOT / "scripts" / "format_manuscript.py"


def load_formatter():
    spec = importlib.util.spec_from_file_location("baker_formatter_tests", FORMATTER_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load formatter: {FORMATTER_PATH}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


FORMATTER = load_formatter()


def build_manuscript(path: Path, paragraph_count: int = 120) -> None:
    doc = Document()
    doc.add_paragraph("A deterministic manuscript formatting test", style="Title")
    doc.add_paragraph("Results", style="Heading 1")
    for index in range(paragraph_count - 3):
        doc.add_paragraph(
            f"Result paragraph {index} reports a preserved quantitative value of {index}."
        )
    doc.add_paragraph("Methods", style="Heading 1")
    doc.save(path)


class FormatManuscriptTests(unittest.TestCase):
    def test_nature_communications_works_without_bundled_template(self) -> None:
        """The public release must not require redistribution of the DOCX asset."""
        with tempfile.TemporaryDirectory(prefix="baker_formatter_test_") as temp_dir:
            workdir = Path(temp_dir)
            source = workdir / "source.docx"
            output = workdir / "formatted.docx"
            report_path = workdir / "audit.json"
            build_manuscript(source)

            with (
                mock.patch.object(
                    FORMATTER,
                    "DEFAULT_NCOMMS_TEMPLATE",
                    workdir / "not-bundled.docx",
                ),
                mock.patch("builtins.print"),
            ):
                return_code = FORMATTER.main(
                    [
                        str(source),
                        "--journal",
                        "nature-communications",
                        "--output",
                        str(output),
                        "--report",
                        str(report_path),
                    ]
                )

            self.assertEqual(return_code, 0)
            self.assertTrue(output.is_file())
            report = json.loads(report_path.read_text(encoding="utf-8"))
            self.assertNotIn("template", report)

    def test_format_document_caches_top_level_paragraphs(self) -> None:
        """Guard against repeated Document.paragraphs reconstruction (O(P^2))."""
        with tempfile.TemporaryDirectory(prefix="baker_formatter_test_") as temp_dir:
            workdir = Path(temp_dir)
            source = workdir / "source.docx"
            output = workdir / "formatted.docx"
            report_path = workdir / "audit.json"
            build_manuscript(source)

            original_getter = DocumentObject.paragraphs.fget
            access_count = 0

            def counted_paragraphs(document):
                nonlocal access_count
                access_count += 1
                return original_getter(document)

            with mock.patch.object(
                DocumentObject, "paragraphs", new=property(counted_paragraphs)
            ):
                report = FORMATTER.format_document(
                    input_path=source,
                    output_path=output,
                    report_path=report_path,
                    profile=FORMATTER.get_profile("nature", "article"),
                    journal="nature",
                    article_type="article",
                    stage="initial",
                    template_path=None,
                    preserve_header_footer=False,
                    audit_only=False,
                )

            self.assertEqual(access_count, 1)
            self.assertEqual(report["formatted_paragraphs"], 120)
            self.assertTrue(output.is_file())
            self.assertTrue(FORMATTER.validate_docx(output) is None)
            persisted_report = json.loads(report_path.read_text(encoding="utf-8"))
            self.assertEqual(persisted_report["output_sha256"], report["output_sha256"])


if __name__ == "__main__":
    unittest.main()
